#IMPORTANTE
#//div/div/div[2]/div[2]/div[1]/div[2]/div[1]/p[@class="ui-promotions-pill ui-pb-highlight-wrapper coupon"]
import json
import re
import requests
import unidecode
import scrapy
import requests
from docx import Document
import pandas
from datetime import datetime

start_row = 20  
end_row = 33
num_rows = end_row - start_row

db = pandas.read_excel("politica-promo.xlsx", engine='openpyxl')

db.columns = ['PRODUTO', 'SITE', 'COLUNA3', 'CLÁSSICO ML', 'COLUNA5', 'PREMIUM ML', 'COLUNA7', 'MARKETPLACES', 'COLUNA9']

df = pandas.read_excel("GESTÃO DE AÇÕES E-COMMERCE.xlsx", usecols='C:O', skiprows=start_row, nrows=num_rows, engine='openpyxl', sheet_name="POLÍTICA COMERCIAL Set24")

df.columns = ['PRODUTO', 'inutil1', 'SITE', 'COLUNA3','inutil2', 'CLÁSSICO ML', 'COLUNA5','inutil3', 'PREMIUM ML', 'COLUNA7','inutil4', 'MARKETPLACES', 'COLUNA9']

for index, i in df.iterrows():
    if i['PRODUTO'] == "FONTE 40A":
        fonte40Marketplace = round(i['COLUNA3'], 2);
        fonte40Classico = round(i['COLUNA5'], 2);
        fonte40Premium = round(i['COLUNA7'], 2);
        fonte40PremiumPrice = round(i['PREMIUM ML'], 2);
        fonte40ClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte40Marketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 60A":
        fonte60Marketplace = round(i['COLUNA3'], 2);
        fonte60Classico = round(i['COLUNA5'], 2);
        fonte60Premium = round(i['COLUNA7'], 2);
        fonte60PremiumPrice = round(i['PREMIUM ML'], 2);
        fonte60ClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte60Marketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 60A LITE":
        fonte60liteMarketplace = round(i['COLUNA3'], 2);
        fonte60liteClassico = round(i['COLUNA5'], 2);
        fonte60litePremium = round(i['COLUNA7'], 2);
        fonte60litePremiumPrice = round(i['PREMIUM ML'], 2);
        fonte60liteClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte60liteMarketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 70A":
        fonte70Marketplace = round(i['COLUNA3'], 2);
        fonte70Classico = round(i['COLUNA5'], 2);
        fonte70Premium = round(i['COLUNA7'], 2);
        fonte70PremiumPrice = round(i['PREMIUM ML'], 2);
        fonte70ClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte70Marketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 70A LITE":
        fonte70liteMarketplace = round(i['COLUNA3'], 2);
        fonte70liteClassico = round(i['COLUNA5'], 2);
        fonte70litePremium = round(i['COLUNA7'], 2);
        fonte70litePremiumPrice = round(i['PREMIUM ML'], 2);
        fonte70liteClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte70liteMarketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 90 BOB":
        fonte90bobMarketplace = round(i['COLUNA3'], 2);
        fonte90bobClassico = round(i['COLUNA5'], 2);
        fonte90bobPremium = round(i['COLUNA7'], 2);
        fonte90bobPremiumPrice = round(i['PREMIUM ML'], 2);
        fonte90bobClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte90bobMarketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 120 BOB":
        fonte120bobMarketplace = round(i['COLUNA3'], 2);
        fonte120bobClassico = round(i['COLUNA5'], 2);
        fonte120bobPremium = round(i['COLUNA7'], 2);
        fonte120bobPremiumPrice = round(i['PREMIUM ML'], 2);
        fonte120bobClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte120bobMarketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 120A LITE":
        fonte120liteMarketplace = round(i['COLUNA3'], 2);
        fonte120liteClassico = round(i['COLUNA5'], 2);
        fonte120litePremium = round(i['COLUNA7'], 2);
        fonte120litePremiumPrice = round(i['PREMIUM ML'], 2);
        fonte120liteClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte120liteMarketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 120A":
        fonte120Marketplace = round(i['COLUNA3'], 2);
        fonte120Classico = round(i['COLUNA5'], 2);
        fonte120Premium = round(i['COLUNA7'], 2);
        fonte120PremiumPrice = round(i['PREMIUM ML'], 2);
        fonte120ClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte120Marketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 200 BOB":
        fonte200bobMarketplace = round(i['COLUNA3'], 2);
        fonte200bobClassico = round(i['COLUNA5'], 2);
        fonte200bobPremium = round(i['COLUNA7'], 2);
        fonte200bobPremiumPrice = round(i['PREMIUM ML'], 2);
        fonte200bobClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte200bobMarketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 200A LITE":
        fonte200liteMarketplace = round(i['COLUNA3'], 2);
        fonte200liteClassico = round(i['COLUNA5'], 2);
        fonte200litePremium = round(i['COLUNA7'], 2);
        fonte200litePremiumPrice = round(i['PREMIUM ML'], 2);
        fonte200liteClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte200liteMarketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 200 MONO":
        fonte200monoMarketplace = round(i['COLUNA3'], 2);
        fonte200monoClassico = round(i['COLUNA5'], 2);
        fonte200monoPremium = round(i['COLUNA7'], 2);
        fonte200monoPremiumPrice = round(i['PREMIUM ML'], 2);
        fonte200monoClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte200monoMarketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 200A":
        fonte200Marketplace = round(i['COLUNA3'], 2);
        fonte200Classico = round(i['COLUNA5'], 2);
        fonte200Premium = round(i['COLUNA7'], 2);
        fonte200PremiumPrice = round(i['PREMIUM ML'], 2);
        fonte200ClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte200Marketplaceprice = round(i['SITE'], 2);
        
for index, i in db.iterrows():
    if i['PRODUTO'] == "FONTE 40A":
        fonte40Marketplace = round(i['COLUNA3'], 2);
        fonte40Classico = round(i['COLUNA5'], 2);
        fonte40Premium = round(i['COLUNA7'], 2);
        fonte40PremiumPrice = round(i['PREMIUM ML'], 2);
        fonte40ClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte40Marketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 60A":
        fonte60Marketplace = round(i['COLUNA3'], 2);
        fonte60Classico = round(i['COLUNA5'], 2);
        fonte60Premium = round(i['COLUNA7'], 2);
        fonte60PremiumPrice = round(i['PREMIUM ML'], 2);
        fonte60ClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte60Marketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 60A LITE":
        fonte60liteMarketplace = round(i['COLUNA3'], 2);
        fonte60liteClassico = round(i['COLUNA5'], 2);
        fonte60litePremium = round(i['COLUNA7'], 2);
        fonte60litePremiumPrice = round(i['PREMIUM ML'], 2);
        fonte60liteClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte60liteMarketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 70A":
        fonte70Marketplace = round(i['COLUNA3'], 2);
        fonte70Classico = round(i['COLUNA5'], 2);
        fonte70Premium = round(i['COLUNA7'], 2);
        fonte70PremiumPrice = round(i['PREMIUM ML'], 2);
        fonte70ClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte70Marketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 70A LITE":
        fonte70liteMarketplace = round(i['COLUNA3'], 2);
        fonte70liteClassico = round(i['COLUNA5'], 2);
        fonte70litePremium = round(i['COLUNA7'], 2);
        fonte70litePremiumPrice = round(i['PREMIUM ML'], 2);
        fonte70liteClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte70liteMarketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 90 BOB":
        fonte90bobMarketplace = round(i['COLUNA3'], 2);
        fonte90bobClassico = round(i['COLUNA5'], 2);
        fonte90bobPremium = round(i['COLUNA7'], 2);
        fonte90bobPremiumPrice = round(i['PREMIUM ML'], 2);
        fonte90bobClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte90bobMarketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 120 BOB":
        fonte120bobMarketplace = round(i['COLUNA3'], 2);
        fonte120bobClassico = round(i['COLUNA5'], 2);
        fonte120bobPremium = round(i['COLUNA7'], 2);
        fonte120bobPremiumPrice = round(i['PREMIUM ML'], 2);
        fonte120bobClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte120bobMarketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 120A LITE":
        fonte120liteMarketplace = round(i['COLUNA3'], 2);
        fonte120liteClassico = round(i['COLUNA5'], 2);
        fonte120litePremium = round(i['COLUNA7'], 2);
        fonte120litePremiumPrice = round(i['PREMIUM ML'], 2);
        fonte120liteClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte120liteMarketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 120A":
        fonte120Marketplace = round(i['COLUNA3'], 2);
        fonte120Classico = round(i['COLUNA5'], 2);
        fonte120Premium = round(i['COLUNA7'], 2);
        fonte120PremiumPrice = round(i['PREMIUM ML'], 2);
        fonte120ClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte120Marketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 200 BOB":
        fonte200bobMarketplace = round(i['COLUNA3'], 2);
        fonte200bobClassico = round(i['COLUNA5'], 2);
        fonte200bobPremium = round(i['COLUNA7'], 2);
        fonte200bobPremiumPrice = round(i['PREMIUM ML'], 2);
        fonte200bobClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte200bobMarketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 200A LITE":
        fonte200liteMarketplace = round(i['COLUNA3'], 2);
        fonte200liteClassico = round(i['COLUNA5'], 2);
        fonte200litePremium = round(i['COLUNA7'], 2);
        fonte200litePremiumPrice = round(i['PREMIUM ML'], 2);
        fonte200liteClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte200liteMarketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 200 MONO":
        fonte200monoMarketplace = round(i['COLUNA3'], 2);
        fonte200monoClassico = round(i['COLUNA5'], 2);
        fonte200monoPremium = round(i['COLUNA7'], 2);
        fonte200monoPremiumPrice = round(i['PREMIUM ML'], 2);
        fonte200monoClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte200monoMarketplaceprice = round(i['SITE'], 2);
    elif i['PRODUTO'] == "FONTE 200A":
        fonte200Marketplace = round(i['COLUNA3'], 2);
        fonte200Classico = round(i['COLUNA5'], 2);
        fonte200Premium = round(i['COLUNA7'], 2);
        fonte200PremiumPrice = round(i['PREMIUM ML'], 2);
        fonte200ClassicoPrice = round(i['CLÁSSICO ML'], 2);
        fonte200Marketplaceprice = round(i['SITE'], 2);

# if os.path.exists("dados_scrapy.docx"):
#     doc = Document("dados_scrapy.docx")
# else:

doc = Document()

def extract_price(response):
  price_selectors = [
      '//*[@id="price"]/div/div[1]/div[1]/span[1]/span/span[2]/text()',
      '//html/body/main/div[2]/div[5]/div/div[1]/div/div[1]/div/div[@class="ui-pdp-container__row ui-pdp-container__row--price"]/div/div[1]/div[1]/span/span/span[2]/text()',
      '//*[@id="ui-pdp-main-container"]/div[1]/div/div[1]/div[2]/div[3]/div[1]/div[1]/span/span/span[2]/text()',
      '//*[@id="ui-pdp-main-container"]/div[1]/div/div[1]/div[2]/div[2]/div[1]/div[1]/span[1]/span/span[2]/text()'
  ]
  
  for selector in price_selectors:
    price = response.xpath(selector).get()
    if price:
      price = price.replace('.', '')
      decimal_selector = selector.replace("span[2]/text()", "") + 'span[@class="andes-money-amount__cents andes-money-amount__cents--superscript-36"]/text()'
      price_decimal = response.xpath(decimal_selector).get()
      
      if price_decimal:
        return float(f"{price}.{price_decimal}")
      else:
        try:
          return float(price)
        except ValueError:
          pass

  return None  

def extract_real_price(response):
    price_selectors = [
        '//span/s/span[@class="andes-money-amount__fraction"]/text()',
        './/div/div/div[3]/div[3]/div[1]/div[1]/div/div/s/span[@class="andes-money-amount__fraction"]/text()'
    ]

    for selector in price_selectors:
        price = response.xpath(selector).get()
        if price:
            price = price.replace('.', '')
            decimal_selector = selector.replace('span[@class="andes-money-amount__fraction"]/text()', '') + 'span[@class="andes-money-amount__cents andes-money-amount__cents--superscript-16"]/text()'
            price_decimal = response.xpath(decimal_selector).get()
            
            if price_decimal:
                return float(f"{price}.{price_decimal}")
            else:
                try:
                    return float(price)
                except ValueError:
                    pass

    return None  

def extract_real_price_marketplace(response):
    price_selectors = [

        './/div/div/div[3]/div[3]/div[1]/div[1]/div/div/s/span[@class="andes-money-amount__fraction"]/text()'
    ]

    for selector in price_selectors:
        price = response.xpath(selector).get()
        if price:
            price = price.replace('.', '')
            decimal_selector = selector.replace('span[@class="andes-money-amount__fraction"]/text()', '') + 'span[@class="andes-money-amount__cents andes-money-amount__cents--superscript-16"]/text()'
            price_decimal = response.xpath(decimal_selector).get()
            
            if price_decimal:
                return float(f"{price}.{price_decimal}")
            else:
                try:
                    return float(price)
                except ValueError:
                    pass

    return None  

def extract_price_new(response):
  price_selectors = [
      './/div/div/div[2]/div[2]/div[1]/div[1]/div/div/div/span[1]/span[@class="andes-money-amount__fraction"]/text()',
      './/div[1]/div[1]/div/div/div/span[1]/span[@class="andes-money-amount__fraction"]/text()',
      './/span[@class="andes-money-amount ui-pdp-price__part andes-money-amount--cents-superscript andes-money-amount--compact"]/span[@class="andes-money-amount__fraction"]/text()'
  ]
  
  for selector in price_selectors:
    price = response.xpath(selector).get()
    if price:
      price = price.replace('.', '')
      decimal_selector = selector.replace('span[@class="andes-money-amount__fraction"]/text()', '') + 'span[@class="andes-money-amount__cents andes-money-amount__cents--superscript-24"]/text()'
      price_decimal = response.xpath(decimal_selector).get()
      if not price_decimal:  
        decimal_selector = selector.replace('span[@class="andes-money-amount__fraction"]/text()', '') + 'span[@class="andes-money-amount__cents andes-money-amount__cents--superscript-20"]/text()'
        price_decimal = response.xpath(decimal_selector).get()
      
      if price_decimal:
        return float(f"{price}.{price_decimal}")
      else:
        try:
          return float(price)
        except ValueError:
          pass

  return None  


class MlSpider(scrapy.Spider):
    option_selected = ""
    option_selected_new = ""
    name = 'ml'
    start_urls = ["https://lista.mercadolivre.com.br/fonte-jfa"]
    
    def __init__(self, palavra=None, cookie=None, *args, **kwargs):
        super(MlSpider, self).__init__(*args, **kwargs)
        with open('cookies.json', 'r') as f:
            self.cookies = json.load(f)
        self.palavra = palavra
        self.cookie = cookie 
    
    def parse(self, response, **kwargs):
        self.option_selected = self.palavra
        self.option_selected_new = self.palavra
        search = ""
        if self.option_selected == "FONTE 40A":
            search = "fonte storm 40a"
        if self.option_selected == "FONTE 60A LITE":
            search = "fonte lite 60a"
        elif self.option_selected == "FONTE 60A":
            search = "fonte storm 60a"
        if self.option_selected == "FONTE 70A LITE":
            search = "fonte lite 70a"
        elif self.option_selected == "FONTE 70A":
            search = "fonte storm 70a"
        elif self.option_selected == "FONTE 90 BOB":
            search = "fonte bob 90a"
        elif self.option_selected == "FONTE 120A":
            search = "fonte storm 120a"
        elif self.option_selected == "FONTE 120A LITE":
            search = "fonte lite 120a"
        elif self.option_selected == "FONTE 120 BOB":
            search = "fonte bob 120a"
        elif self.option_selected == "FONTE 200A":
            search = "fonte storm 200a"
        elif self.option_selected == "FONTE 200A LITE":
            search = "fonte lite 200a"
        elif self.option_selected == "FONTE 200 BOB":
            search = "fonte bob 200a"
        elif self.option_selected == "FONTE 200 MONO":
            search = "fonte storm 200a mono"
        #search = search.replace(" ", "%20")
        
        search_catalog = ""
        if self.option_selected == "FONTE 40A":
            search_catalog = "fonte 40a"
        if self.option_selected == "FONTE 60A LITE":
            search_catalog = "fonte 60a"
        elif self.option_selected == "FONTE 60A":
            search_catalog = "fonte 60a"
        if self.option_selected == "FONTE 70A LITE":
            search_catalog = "fonte 70a"
        elif self.option_selected == "FONTE 70A":
            search_catalog = "fonte 70a"
        elif self.option_selected == "FONTE 90 BOB":
            search_catalog = "fonte 90a"
        elif self.option_selected == "FONTE 120A":
            search_catalog = "fonte 120a"
        elif self.option_selected == "FONTE 120A LITE":
            search_catalog = "fonte 120a"
        elif self.option_selected == "FONTE 120 BOB":
            search_catalog = "fonte 120a"
        elif self.option_selected == "FONTE 200A":
            search_catalog = "fonte 200a"
        elif self.option_selected == "FONTE 200A LITE":
            search_catalog = "fonte 200a"
        elif self.option_selected == "FONTE 200 BOB":
            search_catalog = "fonte 200a"
        elif self.option_selected == "FONTE 200 MONO":
            search_catalog = "fonte 200a mono"
        search_catalog = search_catalog.replace(" ", "%20")
        
        # yield scrapy.Request(url=f"https://lista.mercadolivre.com.br/acessorios-veiculos/{search}_OrderId_PRICE_NoIndex_True", callback=self.parse_all)BRAND_22292586
        # yield scrapy.Request(url=f"https://lista.mercadolivre.com.br/acessorios-veiculos/{search}_Frete_Full_OrderId_PRICE_NoIndex_True", callback=self.parse_all)
        yield scrapy.Request(url=f"https://lista.mercadolivre.com.br/acessorios-veiculos/{search}_Frete_Full_OrderId_PRICE_BRAND_2466336_NoIndex_True", callback=self.parse_all)
        yield scrapy.Request(url=f"https://lista.mercadolivre.com.br/acessorios-veiculos/{search}_OrderId_PRICE_BRAND_2466336_NoIndex_True", callback=self.parse_all)
        yield scrapy.Request(url=f"https://lista.mercadolivre.com.br/acessorios-veiculos/{search}_Frete_Full_OrderId_PRICE_BRAND_22292586_NoIndex_True", callback=self.parse_all)
        yield scrapy.Request(url=f"https://lista.mercadolivre.com.br/acessorios-veiculos/{search}_OrderId_PRICE_BRAND_22292586_NoIndex_True", callback=self.parse_all)
        
    
    def parse_all(self, response):
        
        for item in response.xpath('//div/div[3]/section/ol/li[@class="ui-search-layout__item shops__layout-item ui-search-layout__stack"]'):
            new_name = item.xpath('.//h2[@class="ui-search-item__title"]/text()').get()
            if not new_name:
                new_name = item.xpath('.//h2[@class="ui-search-item__title ui-search-item__group__element"]/a/text()').get()
            name = new_name
            price = extract_price_new(response=item)
            if not price:
                print(response.url)
            cupom = ""
            if item.xpath('.//div/div/div[2]/div[2]/div[1]/div[2]/div[1]/p[@class="ui-promotions-pill ui-pb-highlight-wrapper coupon"]').get():
                cupom = item.xpath('.//div/div/div[2]/div[2]/div[1]/div[2]/div[1]/p[@class="ui-promotions-pill ui-pb-highlight-wrapper coupon"]/span/span/span/text()').get().replace("OFF", "")
                if "%" in cupom and price:
                    cupom = int(re.findall(r'(\d+)%', cupom)[0])
                    cupom = f"Cupom: %{cupom} - {round(price - (price *( cupom / 100)), 2)}"
                elif "R$" in cupom and price:
                    cupom = int(re.findall(r'R\$\s?(\d+,\d+|\d+)', cupom)[0])
                    cupom = f"Cupom: R${cupom} - {round(price - cupom, 2)}"
            # if item.xpath('.//ul[@class="ui-search-winner-alternatives ui-search-winner-alternatives__container--top-space"]/li[@class="ui-search-winner-alternatives__item"]/a/div').get():
            #     cupom += " Mais de um item"
            loja = ""
            listing_type = "Not Found"
            if item.xpath('.//span[@class="ui-search-item__group__element ui-search-installments ui-search-color--BLACK"]').get():
                listing_type = "Clássico"
            elif item.xpath('.//span[@class="ui-search-item__group__element ui-search-installments ui-search-color--LIGHT_GREEN"]').get():
                listing_type = "Premium"
            url = item.xpath('.//div/div/div[2]/div[1]/a[@class="ui-search-item__group__element ui-search-link__title-card ui-search-link"]/@href').get()
            if not new_name:
                print(response.url)
                continue
            new_name = unidecode.unidecode(new_name.lower())
            if not url:
                url = item.xpath('.//a[@class="ui-search-item__group__element ui-search-link__title-card ui-search-link"]/@href').get()
            if not url:
                url = item.xpath('.//h2[@class="ui-search-item__title ui-search-item__group__element"]/a/@href').get()
            if "taramps" in new_name or "stetson" in new_name or "usina" in new_name or "controle" in new_name:
                continue
            if self.option_selected == "FONTE 40A":     
                if "bob" not in new_name and "lite" not in new_name and "light" not in new_name  and "controle" not in new_name and "usina" not in new_name and ("jfa" in new_name or "fonte carregador" in new_name or "fonte automotiva" in new_name or "fonte e carregador" in new_name or "carregador de baterias" in new_name):
                    if "40a" in new_name or "40" in new_name or "40 amperes" in new_name or "40amperes" in new_name or "36a" in new_name or "36" in new_name or "36 amperes" in new_name or "36amperes" in new_name:
                        # if new_name == "Fonte Automotiva Jfa Storm Lite 40a Bivolt Carregador":
                        if listing_type == "Clássico" and price and cupom == "":
                            if self.option_selected == "FONTE 40A" and price >= fonte40Classico:
                                continue;
                        elif listing_type == "Premium" and price and cupom == "":
                            if self.option_selected == "FONTE 40A" and price >= fonte40Premium:
                                continue;
                        yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})
                        # yield scrapy.Request(url=url.split('?')[0] + '/s', callback=self.get_catalog) 
                        yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-40a-jfa_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.bestonline.com.br/fonte-jfa-40a_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-jfa-40a_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.renovonline.com.br/fonte-jfa-40a_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-jfa-40a_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                        
            elif self.option_selected == "FONTE 60A":
                    
                if "bob" not in new_name and "lite" not in new_name and "light" not in new_name  and "controle" not in new_name and "usina" not in new_name and ("jfa" in new_name or "fonte carregador" in new_name or "fonte automotiva" in new_name or "fonte e carregador" in new_name or "carregador de baterias" in new_name):
                    if "60a" in new_name or "60" in new_name or "60 amperes" in new_name or "60amperes" in new_name or "60 a" in new_name:
                        if listing_type == "Clássico" and price and cupom == "":
                            if self.option_selected == "FONTE 60A" and price >= fonte60Classico:
                                continue;
                        elif listing_type == "Premium" and price and cupom == "":
                            if self.option_selected == "FONTE 60A" and price >= fonte60Premium:
                                continue;
                        yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})
                        # yield scrapy.Request(url=url.split('?')[0] + '/s', callback=self.get_catalog) 
                        yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-60a-jfa_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.bestonline.com.br/fonte-jfa-60a_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-jfa-60a_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.renovonline.com.br/fonte-jfa-60a_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-jfa-60a_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})

                        
            elif self.option_selected == "FONTE 60A LITE":
                if "bob" not in new_name and ("lite" in new_name or "light" in new_name) and "controle" not in new_name and "usina" not in new_name and ("jfa" in new_name or "fonte carregador" in new_name or "fonte automotiva" in new_name or "fonte e carregador" in new_name or "carregador de baterias" in new_name):
                    if "60a" in new_name or "60" in new_name or "60 amperes" in new_name or "60amperes" in new_name or "60 a" in new_name:
                        if listing_type == "Clássico" and price and cupom == "":
                            if self.option_selected == "FONTE 60A LITE" and price >= fonte60liteClassico:
                                continue;
                        elif listing_type == "Premium" and price and cupom == "":
                            if self.option_selected == "FONTE 60A LITE" and price >= fonte60litePremium:
                                continue;
                        yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})
                        # yield scrapy.Request(url=url.split('?')[0] + '/s', callback=self.get_catalog) 
                        yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-60a-lite-jfa_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.bestonline.com.br/fonte-jfa-60a-lite_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-jfa-60a-lite_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.renovonline.com.br/fonte-jfa-60a-lite_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-jfa-60a-lite_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})

                        
            elif self.option_selected == "FONTE 70A":
                
                if "bob" not in new_name and "lite" not in new_name and "light" not in new_name  and "controle" not in new_name and "usina" not in new_name and ("jfa" in new_name or "fonte carregador" in new_name or "fonte automotiva" in new_name or "fonte e carregador" in new_name or "carregador de baterias" in new_name):
                    if "70a" in new_name or "70" in new_name or "70 amperes" in new_name or "70amperes" in new_name or "70 a" in new_name:
                        if listing_type == "Clássico" and price and cupom == "":
                            if self.option_selected == "FONTE 70A" and price >= fonte70Classico:
                                continue;
                        elif listing_type == "Premium" and price and cupom == "":
                            if self.option_selected == "FONTE 70A" and price >= fonte70Premium:
                                continue;
                        yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})
                        # yield scrapy.Request(url=url.split('?')[0] + '/s', callback=self.get_catalog) 
                        yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-70a-jfa_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.bestonline.com.br/fonte-jfa-70a_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-jfa-70a_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.renovonline.com.br/fonte-jfa-70a_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-jfa-70a_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})

                        
                        
            elif self.option_selected == "FONTE 70A LITE":
                
                if "bob" not in new_name and  ("lite" in new_name or "light" in new_name) and "controle" not in new_name and "usina" not in new_name and ("jfa" in new_name or "fonte carregador" in new_name or "fonte automotiva" in new_name or "fonte e carregador" in new_name or "carregador de baterias" in new_name):
                    if "70a" in new_name or "70" in new_name or "70 amperes" in new_name or "70amperes" in new_name or "70 a" in new_name:
                        if listing_type == "Clássico" and price and cupom == "":
                            if self.option_selected == "FONTE 70A LITE" and price >= fonte70liteClassico:
                                continue;
                        elif listing_type == "Premium" and price and cupom == "":
                            if self.option_selected == "FONTE 70A LITE" and price >= fonte70litePremium:
                                continue;
                        yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})
                        # yield scrapy.Request(url=url.split('?')[0] + '/s', callback=self.get_catalog) 
                        yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-70a-lite-jfa_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.bestonline.com.br/fonte-jfa-70a-lite_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-jfa-70a-lite_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.renovonline.com.br/fonte-jfa-70a-lite_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-jfa-70a-lite_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})

                        
                        
            elif self.option_selected == "FONTE 90 BOB":
                
                if "bob" in new_name and "lite" not in new_name and "light" not in new_name  and "controle" not in new_name and "usina" not in new_name and ("jfa" in new_name or "fonte carregador" in new_name or "fonte automotiva" in new_name or "fonte e carregador" in new_name or "carregador de baterias" in new_name):
                    if "90a" in new_name or "90" in new_name or "90 amperes" in new_name or "90amperes" in new_name or "90 a" in new_name:
                        if listing_type == "Clássico" and price and cupom == "":
                            if self.option_selected == "FONTE 90 BOB" and price >= fonte90bobClassico:
                                continue;
                        elif listing_type == "Premium" and price and cupom == "":
                            if self.option_selected == "FONTE 90 BOB" and price >= fonte90bobPremium:
                                continue;
                        yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})
                        # yield scrapy.Request(url=url.split('?')[0] + '/s', callback=self.get_catalog) 
                        yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-90a-bob-jfa_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.bestonline.com.br/fonte-jfa-90a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-jfa-90a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.renovonline.com.br/fonte-jfa-90a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-jfa-90a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})

                        
                        
            elif self.option_selected == "FONTE 120A":
                
                if "bob" not in new_name and "lite" not in new_name and "light" not in new_name  and "controle" not in new_name and "usina" not in new_name and ("jfa" in new_name or "fonte carregador" in new_name or "fonte automotiva" in new_name or "fonte e carregador" in new_name or "carregador de baterias" in new_name):
                    if "120a" in new_name or "120" in new_name or "120 amperes" in new_name or "120amperes" in new_name or "120 a" in new_name:
                        if listing_type == "Clássico" and price and cupom == "":
                            if self.option_selected == "FONTE 120A" and price >= fonte120Classico:
                                continue;
                        elif listing_type == "Premium" and price and cupom == "":
                            if self.option_selected == "FONTE 120A" and price >= fonte120Premium:
                                continue;
                        yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price,'listing_type': listing_type, 'cupom': cupom})
                        # yield scrapy.Request(url=url.split('?')[0] + '/s', callback=self.get_catalog) 
                        yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-120a-jfa_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.bestonline.com.br/fonte-jfa-120a_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-jfa-120a_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.renovonline.com.br/fonte-jfa-120a_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-jfa-120a_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})

                        
                        
            elif self.option_selected == "FONTE 120A LITE":
                
                if "bob" not in new_name and  ("lite" in new_name or "light" in new_name) and "controle" not in new_name and "usina" not in new_name and ("jfa" in new_name or "fonte carregador" in new_name or "fonte automotiva" in new_name or "fonte e carregador" in new_name or "carregador de baterias" in new_name):
                    if "120a" in new_name or "120" in new_name or "120 amperes" in new_name or "120amperes" in new_name or "120 a" in new_name:
                        if listing_type == "Clássico" and price and cupom == "":
                            if self.option_selected == "FONTE 120A LITE" and price >= fonte120liteClassico:
                                continue;
                        elif listing_type == "Premium" and price and cupom == "":
                            if self.option_selected == "FONTE 120A LITE" and price >= fonte120litePremium:
                                continue;
                        yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})
                        # yield scrapy.Request(url=url.split('?')[0] + '/s', callback=self.get_catalog) 
                        yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-120a-lite-jfa_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.bestonline.com.br/fonte-jfa-120a-lite_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-jfa-120a-lite_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.renovonline.com.br/fonte-jfa-120a-lite_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-jfa-120a-lite_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})

                        
                        
            elif self.option_selected == "FONTE 120 BOB":
                
                if "bob" in new_name and "lite" not in new_name and "light" not in new_name  and "controle" not in new_name and "usina" not in new_name and ("jfa" in new_name or "fonte carregador" in new_name or "fonte automotiva" in new_name or "fonte e carregador" in new_name or "carregador de baterias" in new_name):
                    if "120a" in new_name or "120" in new_name or "120 amperes" in new_name or "120amperes" in new_name or "120 a" in new_name:
                        if listing_type == "Clássico" and price and cupom == "":
                            if self.option_selected == "FONTE 120 BOB" and price >= fonte120bobClassico:
                                continue;
                        elif listing_type == "Premium" and price and cupom == "":
                            if self.option_selected == "FONTE 120 BOB" and price >= fonte120bobPremium:
                                continue;
                        yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price,'listing_type': listing_type, 'cupom': cupom})
                        # yield scrapy.Request(url=url.split('?')[0] + '/s', callback=self.get_catalog) 
                        yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-120a-bob-jfa_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.bestonline.com.br/fonte-jfa-120a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-jfa-120a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.renovonline.com.br/fonte-jfa-120a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-jfa-120a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})

                        
                        
            elif self.option_selected == "FONTE 200A":
                
                if "bob" not in new_name and "lite" not in new_name and "light" not in new_name  and "controle" not in new_name and 'mono' not in new_name and 'mono' not in new_name and 'monovolt' not in new_name:
                    if "200a" in new_name or "200" in new_name or "200 amperes" in new_name or "200amperes" in new_name or "200 a" in new_name:
                        if listing_type == "Clássico" and price and cupom == "":
                            if self.option_selected == "FONTE 200A" and price >= fonte200Classico:
                                continue;
                        elif listing_type == "Premium" and price and cupom == "":
                            if self.option_selected == "FONTE 200A" and price >= fonte200Premium:
                                continue;
                        yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})
                        # yield scrapy.Request(url=url.split('?')[0] + '/s', callback=self.get_catalog) 
                        yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-200a-jfa_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.bestonline.com.br/fonte-jfa-200a_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-jfa-200a_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.renovonline.com.br/fonte-jfa-200a_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-jfa-200a_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})

                        
                        
            elif self.option_selected == "FONTE 200 MONO":
                
                if "bob" not in new_name and "lite" not in new_name and "light" not in new_name  and "controle" not in new_name and ("mono" in new_name or "220v" in new_name or "monovolt" in new_name):
                    if "200a" in new_name or "200" in new_name or "200 amperes" in new_name or "200amperes" in new_name or "200 a" in new_name:
                        if listing_type == "Clássico" and price and cupom == "":
                            if self.option_selected == "FONTE 200 MONO" and price >= fonte200monoClassico:
                                continue;
                        elif listing_type == "Premium" and price and cupom == "":
                            if self.option_selected == "FONTE 200 MONO" and price >= fonte200monoPremium:
                                continue;
                        yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price,'listing_type': listing_type, 'cupom': cupom})
                        # yield scrapy.Request(url=url.split('?')[0] + '/s', callback=self.get_catalog) 
                        yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-200a-mono-jfa_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.bestonline.com.br/fonte-jfa-200a-mono_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-jfa-200a-mono_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.renovonline.com.br/fonte-jfa-200a-mono_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-jfa-200a-mono_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})

                        
                        
            elif self.option_selected == "FONTE 200A LITE":
                if "bob" not in new_name and  ("lite" in new_name or "light" in new_name) and "controle" not in new_name and 'mono' not in new_name and 'monovolt' not in new_name:
                    if "200a" in new_name or "200" in new_name or "200 amperes" in new_name or "200amperes" in new_name or "200 a" in new_name:
                        if listing_type == "Clássico" and price and cupom == "":
                            if self.option_selected == "FONTE 200A LITE" and price >= fonte200liteClassico:
                                continue;
                        elif listing_type == "Premium" and price and cupom == "":
                            if self.option_selected == "FONTE 200A LITE" and price >= fonte200litePremium:
                                continue;
                        yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})
                        # yield scrapy.Request(url=url.split('?')[0] + '/s', callback=self.get_catalog) 
                        yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-200a-lite-jfa_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.bestonline.com.br/fonte-jfa-200a-lite_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-jfa-200a-lite_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.renovonline.com.br/fonte-jfa-200a-lite_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-jfa-200a-lite_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})

                        
                        
            elif self.option_selected == "FONTE 200 BOB":
                if "bob" in new_name and "lite" not in new_name and "light" not in new_name  and "controle" not in new_name and 'mono' not in new_name and 'mono' not in new_name and 'monovolt' not in new_name and "usina" not in new_name and ("jfa" in new_name or "fonte carregador" in new_name or "fonte automotiva" in new_name or "fonte e carregador" in new_name or "carregador de baterias" in new_name):
                    if "200a" in new_name or "200" in new_name or "200 amperes" in new_name or "200amperes" in new_name or "200 a" in new_name:
                        if listing_type == "Clássico" and price and cupom == "":
                            if self.option_selected == "FONTE 200 BOB" and price >= fonte200bobClassico:
                                continue;
                        elif listing_type == "Premium" and price and cupom == "":
                            if self.option_selected == "FONTE 200 BOB" and price >= fonte200bobPremium:
                                continue;
                        yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})
                        # yield scrapy.Request(url=url.split('?')[0] + '/s', callback=self.get_catalog) 
                        yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-200a-bob-jfa_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.bestonline.com.br/fonte-jfa-200a-bo_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-jfa-200a-bo_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.renovonline.com.br/fonte-jfa-200a-bo_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                        yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-jfa-200a-bo_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})

        if response.xpath('//nav/ul/li/a[@class="andes-pagination__link" and @title="Seguinte"]'):
            next_page = response.xpath('//nav/ul/li/a[@class="andes-pagination__link" and @title="Seguinte"]/@href').get()
            if next_page:
                yield scrapy.Request(url=next_page, callback=self.parse_all)


                
    def parse_product(self, response):

        cupom = response.meta['cupom']
        name = response.meta['name']
        loja = response.meta['loja']
        listing_type = response.meta['listing_type']
        link = response.xpath('//*[@id="ui-pdp-main-container"]/div[1]/div/div[1]/div/div[3]/div[2]/div/div/div/div/a/@href').get()
        loja = response.xpath('//div[1]/div/button[@class="ui-pdp-seller__link-trigger-button non-selectable"]/span[2]/text()').get()
        self.option_selected_new = self.option_selected
        new_price_float = response.meta["price"]
        real_price = extract_real_price(response)
        
        if self.option_selected == "FONTE 40A" and real_price:
            if listing_type == "Premium":
                if real_price >= fonte40PremiumPrice:
                    return
            else:
                if real_price >= fonte40ClassicoPrice:
                    return
        elif self.option_selected == "FONTE 60A LITE" and real_price:
            if listing_type == "Premium":
                if real_price >= fonte60litePremiumPrice:
                    return
            else:
                if real_price >= fonte60liteClassicoPrice:
                    return
        elif self.option_selected == "FONTE 60A" and real_price:
            if listing_type == "Premium":
                if real_price >= fonte60PremiumPrice:
                    return
            else:
                if real_price >= fonte60ClassicoPrice:
                    return
        elif self.option_selected == "FONTE 70A LITE" and real_price:
            if listing_type == "Premium":
                if real_price >= fonte70litePremiumPrice:
                    return
            else:
                if real_price >= fonte70liteClassicoPrice:
                    return
        elif self.option_selected == "FONTE 70A" and real_price:
            if listing_type == "Premium":
                if real_price >= fonte70PremiumPrice:
                    return
            else:
                if real_price >= fonte70ClassicoPrice:
                    return
        elif self.option_selected == "FONTE 90 BOB" and real_price:
            if listing_type == "Premium":
                if real_price >= fonte90bobPremiumPrice:
                    return
            else:
                if real_price >= fonte90bobClassicoPrice:
                    return
        elif self.option_selected == "FONTE 120A" and real_price:
            if listing_type == "Premium":
                if real_price >= fonte120PremiumPrice:
                    return
            else:
                if real_price >= fonte120ClassicoPrice:
                    return
        elif self.option_selected == "FONTE 120A LITE" and real_price:
            if listing_type == "Premium":
                if real_price >= fonte120litePremiumPrice:
                    return
            else:
                if real_price >= fonte120liteClassicoPrice:
                    return
        elif self.option_selected == "FONTE 120 BOB" and real_price:
            if listing_type == "Premium":
                if real_price >= fonte120bobPremiumPrice:
                    return
            else:
                if real_price >= fonte120bobClassicoPrice:
                    return
        elif self.option_selected == "FONTE 200A" and real_price:
            if listing_type == "Premium":
                if real_price >= fonte200PremiumPrice:
                    return
            else:
                if real_price >= fonte200ClassicoPrice:
                    return
        elif self.option_selected == "FONTE 200A LITE" and real_price:
            if listing_type == "Premium":
                if real_price >= fonte200litePremiumPrice:
                    return
            else:
                if real_price >= fonte200liteClassicoPrice:
                    return
        elif self.option_selected == "FONTE 200 BOB" and real_price:
            if listing_type == "Premium":
                if real_price >= fonte200bobPremiumPrice:
                    return
            else:
                if real_price >= fonte200bobClassicoPrice:
                    return
        elif self.option_selected == "FONTE 200 MONO" and real_price:
            if listing_type == "Premium":
                if real_price >= fonte200monoPremiumPrice:
                    return
            else:
                if real_price >= fonte200monoClassicoPrice:
                    return

                
        tipo = listing_type
          
        
        location_url = f'https://www.mercadolivre.com.br/perfil/{loja.replace(" ", "+")}'
        

        for i in response.xpath('//section/div[2]/div/div/div/div[1]/div/table/tbody/tr'):
            if i.xpath('.//th/div[@class="andes-table__header__container"]/text()').get().lower() == "modelo" or i.xpath('.//th/div[@class="andes-table__header__container"]/text()').get().lower() == "linha":
                modelo = i.xpath('.//td/span/text()').get()
                if modelo:
                    modelo = modelo.lower()
                    if self.option_selected:
                        if self.option_selected == "FONTE 200A":
                            if "bob" in modelo:
                                return

        yield scrapy.Request(url=location_url, callback=self.parse_location, meta={'link_cupom': link,'url': response.url, 'name': name, 'price': new_price_float, 'qtde_parcelado': 0, 'price_parcelado': 0, 'loja': loja, 'tipo': tipo, 'cupom': cupom})


    def finish(self, total_price, url, nomeFonte, loja, lugar):
        if self.option_selected_new == "FONTE 40A" and total_price >= fonte40Marketplace:
            return;
        elif self.option_selected_new == "FONTE 60A LITE" and total_price >= fonte60liteMarketplace:
            return;
        elif self.option_selected_new == "FONTE 60A" and total_price >= fonte60Marketplace:
            return;
        elif self.option_selected_new == "FONTE 70A LITE" and total_price >= fonte70liteMarketplace:
            return;
        elif self.option_selected_new == "FONTE 70A" and total_price >= fonte70Marketplace:
            return;
        elif self.option_selected_new == "FONTE 90 BOB" and total_price >= fonte90bobMarketplace:
            return;
        elif self.option_selected_new == "FONTE 120 BOB" and total_price >= fonte120bobMarketplace:
            return;
        elif self.option_selected_new == "FONTE 120A LITE" and total_price >= fonte120liteMarketplace:
            return;
        elif self.option_selected_new == "FONTE 120A" and total_price >= fonte120Marketplace:
            return;
        elif self.option_selected_new == "FONTE 200 BOB" and total_price >= fonte200bobMarketplace:
            return;
        elif self.option_selected_new == "FONTE 200A LITE" and total_price >= fonte200liteMarketplace:
            return;
        elif self.option_selected_new == "FONTE 200 MONO" and total_price >= fonte200monoMarketplace:
            return;
        elif self.option_selected_new == "FONTE 200A" and total_price >= fonte200Marketplace:
            return;
        
        parcelado = self.get_price_previsto("NA")

        doc.add_paragraph(f'Modelo: {self.option_selected_new}')
        doc.add_paragraph(f'URL: {url}')
        doc.add_paragraph(f'Nome: {nomeFonte}')
        doc.add_paragraph(f'Preço: {total_price}')
        doc.add_paragraph(f'Preço Previsto: {parcelado}')
        doc.add_paragraph(f'Loja: {loja}')
        doc.add_paragraph('Tipo: ')
        doc.add_paragraph(f'Lugar: {lugar}')
        doc.add_paragraph(f'Cupom: ')
        doc.add_paragraph("--------------------------------------------------------------------")
        doc.add_paragraph('')
        doc.save(fr"dados/{self.option_selected_new}.docx")
        if url != None:
            yield {
                'url': url,
                'name': nomeFonte,
                'price': total_price,
                'loja': loja,
                'tipo': "",
                'lugar': lugar
            }


    def parse_radicalson(self, response):
        loja = "RADICALSOM"
        lugar = "Artur nogueira, São Paulo."
        for i in response.xpath('//*[@id="root-app"]/div/div[3]/section/ol/li'):
            nomeFonte = i.xpath('.//h2[@class="ui-search-item__title ui-search-item__group__element shops__items-group-details shops__item-title"]/a/text()').get()
            price = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[2]/text()').get()
            cents = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[4]/text()').get()
            url = i.xpath('.//h2[@class="ui-search-item__title ui-search-item__group__element shops__items-group-details shops__item-title"]/a/@href').get()
            if not url:
                url = i.xpath('.//a[@class="ui-search-item__group__element ui-search-link__title-card shops__items-group-details ui-search-link"]/@href').get()
            nomeFonte = nomeFonte.lower()
            nomeFonte = unidecode.unidecode(nomeFonte)
            real_price = extract_real_price_marketplace(response=i)
            if not cents:
                cents = 0
            if price:
                price = price.replace('.', '')
                total_price = float(f"{price}.{cents}")
            if self.option_selected == "FONTE 40A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "40a" in nomeFonte or "40" in nomeFonte or "40 amperes" in nomeFonte or "40amperes" in nomeFonte or "36a" in nomeFonte or "36" in nomeFonte or "36 amperes" in nomeFonte or "36amperes" in nomeFonte:
                        if real_price:
                            if real_price >= fonte40Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 60A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte or "60 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte60Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 60A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte or "60 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte60liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 70A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte or "70 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte70Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 70A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte or "70 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte70liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 90 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "90a" in nomeFonte or "90" in nomeFonte or "90 amperes" in nomeFonte or "90amperes" in nomeFonte or "90 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte90bobMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 120A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte120Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 120A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte120liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 120 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte120bobMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte and '220' not in nomeFonte and '220v' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200 MONO":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and ("mono" in nomeFonte or "220v" in nomeFonte or "monovolt" in nomeFonte):
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200monoMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200bobMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                    
    def parse_lsdistribuidora(self, response):
        loja = "LS DISTRIBUIDORA"
        lugar = "Elísio Medrado, Bahia"
        for i in response.xpath('//*[@id="root-app"]/div/div[3]/section/ol/li'):
            nomeFonte = i.xpath('.//h2[@class="ui-search-item__title ui-search-item__group__element shops__items-group-details shops__item-title"]/a/text()').get()
            price = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[2]/text()').get()
            cents = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[4]/text()').get()
            url = i.xpath('.//h2[@class="ui-search-item__title ui-search-item__group__element shops__items-group-details shops__item-title"]/a/@href').get()
            if not url:
                url = i.xpath('.//a[@class="ui-search-item__group__element ui-search-link__title-card shops__items-group-details ui-search-link"]/@href').get()
            nomeFonte = nomeFonte.lower()
            nomeFonte = unidecode.unidecode(nomeFonte)
            real_price = extract_real_price_marketplace(response=i)
            if not cents:
                cents = 0
            if price:
                price = price.replace('.', '')
                total_price = float(f"{price}.{cents}")
            if self.option_selected == "FONTE 40A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "40a" in nomeFonte or "40" in nomeFonte or "40 amperes" in nomeFonte or "40amperes" in nomeFonte or "36a" in nomeFonte or "36" in nomeFonte or "36 amperes" in nomeFonte or "36amperes" in nomeFonte:
                        if real_price:
                            if real_price >= fonte40Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 60A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte or "60 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte60Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 60A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte or "60 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte60liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 70A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte or "70 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte70Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 70A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte or "70 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte70liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 90 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "90a" in nomeFonte or "90" in nomeFonte or "90 amperes" in nomeFonte or "90amperes" in nomeFonte or "90 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte90bobMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 120A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte120Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 120A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte120liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 120 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte120bobMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte and '220' not in nomeFonte and '220v' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200 MONO":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and ("mono" in nomeFonte or "220v" in nomeFonte or "monovolt" in nomeFonte):
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200monoMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200bobMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)       

    
    def parse_bestonline(self, response):
        loja = "BESTONLINE"
        lugar = "Rosario, Santa Fe."
        for i in response.xpath('//li[@class="ui-search-layout__item shops__layout-item shops__layout-item ui-search-layout__stack"]'):
            nomeFonte = i.xpath('.//h2[@class="ui-search-item__title ui-search-item__group__element shops__items-group-details shops__item-title"]/a/text()').get()
            price = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[2]/text()').get()
            cents = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[4]/text()').get()
            url = i.xpath('.//h2[@class="ui-search-item__title ui-search-item__group__element shops__items-group-details shops__item-title"]/a/@href').get()
            if not url:
                url = i.xpath('.//a[@class="ui-search-item__group__element ui-search-link__title-card shops__items-group-details ui-search-link"]/@href').get()
            nomeFonte = nomeFonte.lower()
            nomeFonte = unidecode.unidecode(nomeFonte)
            real_price = extract_real_price_marketplace(response=i)
            if not cents:
                cents = 0
            if price:
                price = price.replace('.', '')
                total_price = float(f"{price}.{cents}")
            if self.option_selected == "FONTE 40A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "40a" in nomeFonte or "40" in nomeFonte or "40 amperes" in nomeFonte or "40amperes" in nomeFonte or "36a" in nomeFonte or "36" in nomeFonte or "36 amperes" in nomeFonte or "36amperes" in nomeFonte:
                        if real_price:
                            if real_price >= fonte40Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 60A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte or "60 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte60Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 60A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte or "60 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte60liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 70A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte or "70 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte70Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 70A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte or "70 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte70liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 90 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "90a" in nomeFonte or "90" in nomeFonte or "90 amperes" in nomeFonte or "90amperes" in nomeFonte or "90 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte90bobMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 120A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte120Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 120A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte120liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 120 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte120bobMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte and '220' not in nomeFonte and '220v' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200 MONO":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and ("mono" in nomeFonte or "220v" in nomeFonte or "monovolt" in nomeFonte):
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200monoMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200bobMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)      
            
    
    def parse_renovonline(self, response):
        loja = "RENOV ONLINE"
        lugar = "São João da Boa Vista - SP"
        for i in response.xpath('//li[@class="ui-search-layout__item shops__layout-item shops__layout-item ui-search-layout__stack"]'):
            nomeFonte = i.xpath('.//h2[@class="ui-search-item__title ui-search-item__group__element shops__items-group-details shops__item-title"]/a/text()').get()
            price = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[2]/text()').get()
            cents = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[4]/text()').get()
            url = i.xpath('.//h2[@class="ui-search-item__title ui-search-item__group__element shops__items-group-details shops__item-title"]/a/@href').get()
            if not url:
                url = i.xpath('.//a[@class="ui-search-item__group__element ui-search-link__title-card shops__items-group-details ui-search-link"]/@href').get()
            nomeFonte = nomeFonte.lower()
            nomeFonte = unidecode.unidecode(nomeFonte)
            real_price = extract_real_price_marketplace(response=i)
            if not cents:
                cents = 0
            if price:
                price = price.replace('.', '')
                total_price = float(f"{price}.{cents}")
            if self.option_selected == "FONTE 40A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "40a" in nomeFonte or "40" in nomeFonte or "40 amperes" in nomeFonte or "40amperes" in nomeFonte or "36a" in nomeFonte or "36" in nomeFonte or "36 amperes" in nomeFonte or "36amperes" in nomeFonte:
                        if real_price:
                            if real_price >= fonte40Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 60A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte or "60 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte60Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 60A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte or "60 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte60liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 70A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte or "70 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte70Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 70A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte or "70 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte70liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 90 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "90a" in nomeFonte or "90" in nomeFonte or "90 amperes" in nomeFonte or "90amperes" in nomeFonte or "90 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte90bobMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 120A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte120Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 120A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte120liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 120 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte120bobMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte and '220' not in nomeFonte and '220v' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200 MONO":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and ("mono" in nomeFonte or "220v" in nomeFonte or "monovolt" in nomeFonte):
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200monoMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200bobMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)                     
        
    def parse_shoppratico(self, response):
        loja = "SHOPPRATICO"
        lugar = "Sorocaba, São Paulo."
        for i in response.xpath('//li[@class="ui-search-layout__item shops__layout-item shops__layout-item ui-search-layout__stack"]'):
            nomeFonte = i.xpath('.//h2[@class="ui-search-item__title ui-search-item__group__element shops__items-group-details shops__item-title"]/a/text()').get()
            price = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[2]/text()').get()
            cents = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[4]/text()').get()
            url = i.xpath('.//h2[@class="ui-search-item__title ui-search-item__group__element shops__items-group-details shops__item-title"]/a/@href').get()
            if not url:
                url = i.xpath('.//a[@class="ui-search-item__group__element ui-search-link__title-card shops__items-group-details ui-search-link"]/@href').get()
            nomeFonte = nomeFonte.lower()
            nomeFonte = unidecode.unidecode(nomeFonte)
            real_price = extract_real_price_marketplace(response=i)
            if not cents:
                cents = 0
            if price:
                price = price.replace('.', '')
                total_price = float(f"{price}.{cents}")
            if self.option_selected == "FONTE 40A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "40a" in nomeFonte or "40" in nomeFonte or "40 amperes" in nomeFonte or "40amperes" in nomeFonte or "36a" in nomeFonte or "36" in nomeFonte or "36 amperes" in nomeFonte or "36amperes" in nomeFonte:
                        if real_price:
                            if real_price >= fonte40Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 60A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte or "60 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte60Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 60A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte or "60 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte60liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 70A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte or "70 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte70Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 70A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte or "70 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte70liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 90 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "90a" in nomeFonte or "90" in nomeFonte or "90 amperes" in nomeFonte or "90amperes" in nomeFonte or "90 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte90bobMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 120A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte120Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 120A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte120liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 120 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte120bobMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte and '220' not in nomeFonte and '220v' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200Marketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200 MONO":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and ("mono" in nomeFonte or "220v" in nomeFonte or "monovolt" in nomeFonte):
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200monoMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200liteMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            elif self.option_selected == "FONTE 200 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        if real_price:
                            if real_price >= fonte200bobMarketplaceprice:
                                continue
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)            
        
    def get_price_previsto(self, tipo):
        if tipo == "Clássico":
            for index, i in db.iterrows():
                if self.option_selected_new == "FONTE 40A" and i['PRODUTO'] == "FONTE 40A":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 60A" and i['PRODUTO'] == "FONTE 60A":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 60A LITE" and i['PRODUTO'] == "FONTE 60A LITE":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 70A" and i['PRODUTO'] == "FONTE 70A":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 70A LITE" and i['PRODUTO'] == "FONTE 70A LITE":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 90 BOB" and i['PRODUTO'] == "FONTE 90 BOB":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 120 BOB" and i['PRODUTO'] == "FONTE 120 BOB":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 120A LITE" and i['PRODUTO'] == "FONTE 120A LITE":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 120A" and i['PRODUTO'] == "FONTE 120A":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 200 BOB" and i['PRODUTO'] == "FONTE 200 BOB":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 200A LITE" and i['PRODUTO'] == "FONTE 200A LITE":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 200 MONO" and i['PRODUTO'] == "FONTE 200 MONO":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 200A" and i['PRODUTO'] == "FONTE 200A":
                    return round(i['COLUNA5'], 2);
            for index, i in df.iterrows():
                if self.option_selected_new == "FONTE 40A" and i['PRODUTO'] == "FONTE 40A":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 60A" and i['PRODUTO'] == "FONTE 60A":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 60A LITE" and i['PRODUTO'] == "FONTE 60A LITE":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 70A" and i['PRODUTO'] == "FONTE 70A":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 70A LITE" and i['PRODUTO'] == "FONTE 70A LITE":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 90 BOB" and i['PRODUTO'] == "FONTE 90 BOB":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 120 BOB" and i['PRODUTO'] == "FONTE 120 BOB":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 120A LITE" and i['PRODUTO'] == "FONTE 120A LITE":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 120A" and i['PRODUTO'] == "FONTE 120A":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 200 BOB" and i['PRODUTO'] == "FONTE 200 BOB":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 200A LITE" and i['PRODUTO'] == "FONTE 200A LITE":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 200 MONO" and i['PRODUTO'] == "FONTE 200 MONO":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 200A" and i['PRODUTO'] == "FONTE 200A":
                    return round(i['COLUNA5'], 2);
        elif tipo == "Premium":
            for index, i in db.iterrows():
                if self.option_selected_new == "FONTE 40A" and i['PRODUTO'] == "FONTE 40A":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 60A" and i['PRODUTO'] == "FONTE 60A":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 60A LITE" and i['PRODUTO'] == "FONTE 60A LITE":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 70A" and i['PRODUTO'] == "FONTE 70A":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 70A LITE" and i['PRODUTO'] == "FONTE 70A LITE":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 90 BOB" and i['PRODUTO'] == "FONTE 90 BOB":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 120 BOB" and i['PRODUTO'] == "FONTE 120 BOB":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 120A LITE" and i['PRODUTO'] == "FONTE 120A LITE":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 120A" and i['PRODUTO'] == "FONTE 120A":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 200 BOB" and i['PRODUTO'] == "FONTE 200 BOB":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 200A LITE" and i['PRODUTO'] == "FONTE 200A LITE":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 200 MONO" and i['PRODUTO'] == "FONTE 200 MONO":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 200A" and i['PRODUTO'] == "FONTE 200A":
                    return round(i['COLUNA7'], 2);
            for index, i in df.iterrows():
                if self.option_selected_new == "FONTE 40A" and i['PRODUTO'] == "FONTE 40A":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 60A" and i['PRODUTO'] == "FONTE 60A":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 60A LITE" and i['PRODUTO'] == "FONTE 60A LITE":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 70A" and i['PRODUTO'] == "FONTE 70A":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 70A LITE" and i['PRODUTO'] == "FONTE 70A LITE":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 90 BOB" and i['PRODUTO'] == "FONTE 90 BOB":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 120 BOB" and i['PRODUTO'] == "FONTE 120 BOB":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 120A LITE" and i['PRODUTO'] == "FONTE 120A LITE":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 120A" and i['PRODUTO'] == "FONTE 120A":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 200 BOB" and i['PRODUTO'] == "FONTE 200 BOB":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 200A LITE" and i['PRODUTO'] == "FONTE 200A LITE":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 200 MONO" and i['PRODUTO'] == "FONTE 200 MONO":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 200A" and i['PRODUTO'] == "FONTE 200A":
                    return round(i['COLUNA7'], 2);
        elif tipo == "NA":
            for index, i in db.iterrows():
                if self.option_selected_new == "FONTE 40A" and i['PRODUTO'] == "FONTE 40A":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 60A" and i['PRODUTO'] == "FONTE 60A":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 60A LITE" and i['PRODUTO'] == "FONTE 60A LITE":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 70A" and i['PRODUTO'] == "FONTE 70A":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 70A LITE" and i['PRODUTO'] == "FONTE 70A LITE":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 90 BOB" and i['PRODUTO'] == "FONTE 90 BOB":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 120 BOB" and i['PRODUTO'] == "FONTE 120 BOB":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 120A LITE" and i['PRODUTO'] == "FONTE 120A LITE":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 120A" and i['PRODUTO'] == "FONTE 120A":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 200 BOB" and i['PRODUTO'] == "FONTE 200 BOB":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 200A LITE" and i['PRODUTO'] == "FONTE 200A LITE":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 200 MONO" and i['PRODUTO'] == "FONTE 200 MONO":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 200A" and i['PRODUTO'] == "FONTE 200A":
                    return round(i['COLUNA3'], 2);
            for index, i in df.iterrows():
                if self.option_selected_new == "FONTE 40A" and i['PRODUTO'] == "FONTE 40A":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 60A" and i['PRODUTO'] == "FONTE 60A":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 60A LITE" and i['PRODUTO'] == "FONTE 60A LITE":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 70A" and i['PRODUTO'] == "FONTE 70A":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 70A LITE" and i['PRODUTO'] == "FONTE 70A LITE":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 90 BOB" and i['PRODUTO'] == "FONTE 90 BOB":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 120 BOB" and i['PRODUTO'] == "FONTE 120 BOB":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 120A LITE" and i['PRODUTO'] == "FONTE 120A LITE":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 120A" and i['PRODUTO'] == "FONTE 120A":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 200 BOB" and i['PRODUTO'] == "FONTE 200 BOB":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 200A LITE" and i['PRODUTO'] == "FONTE 200A LITE":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 200 MONO" and i['PRODUTO'] == "FONTE 200 MONO":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 200A" and i['PRODUTO'] == "FONTE 200A":
                    return round(i['COLUNA3'], 2);

    def parse_location(self, response):
        name = response.meta['name']
        url = response.meta['url']
        new_price_float = response.meta['price']
        tipo = response.meta['tipo']
        cupom = response.meta['cupom']
        link = response.meta['link_cupom']
        parcelado = self.get_price_previsto(tipo)
        loja = response.meta['loja']
        lugar = response.xpath('//*[@id="profile"]/div/div[2]/div[1]/div[3]/p/text()').get()

        if link != None:
            yield scrapy.Request(url=link, callback=self.get_cupom, meta={'url': url, 'name': name, 'price': new_price_float, 'loja': loja, 'tipo': tipo, 'cupom': cupom, 'lugar': lugar, 'parcelado': parcelado})
            
        if not cupom:
            doc.add_paragraph(f'Modelo: {self.option_selected_new}')
            doc.add_paragraph(f'URL: {url}')
            doc.add_paragraph(f'Nome: {name}')
            doc.add_paragraph(f'Preço: {new_price_float}')
            doc.add_paragraph(f'Preço Previsto: {parcelado}')
            doc.add_paragraph(f'Loja: {loja}')
            doc.add_paragraph(f'Tipo: {tipo}')
            doc.add_paragraph(f'Lugar: {lugar}')
            doc.add_paragraph(f'Cupom: {cupom}')
            doc.add_paragraph("--------------------------------------------------------------------")
            doc.add_paragraph('')
            
            yield {
                'url': url,
                'name': name,
                'price': new_price_float,
                'price_previsto': parcelado,
                'loja': loja,
                'tipo': tipo,
                'lugar': lugar
            }
            doc.save(fr"dados/{self.option_selected_new}.docx")
         
    def get_catalog(self, response):
        for i in response.xpath('//*[@id="buybox-form"]'):
            price = extract_price_new(response=i)
            listing_type = ""
            if "sem juros" in i.xpath('.//p[@class="ui-pdp-family--REGULAR ui-pdp-media__title"]/text()'):
                listing_type = "Premium"
            else:
                listing_type = "Clássico"
            loja = i.xpath('.//button[@class="ui-pdp-seller__link-trigger-button non-selectable"]/span/text()').get()
            if (i.xpath('.//div[@class="ui-pdp-actions__container"]/input[@name="item_id"]/@value')):
                url = response.url.split('/s')[0] + "?pdp_filters=" + i.xpath('.//div[@class="ui-pdp-actions__container"]/input[@name="item_id"]/@value').get()
            else:
                print(response.url)
            if self.option_selected == "FONTE 40A":     
                if listing_type == "Clássico" and price:
                    if self.option_selected == "FONTE 40A" and price >= fonte40Classico:
                        continue;
                elif listing_type == "Premium" and price:
                    if self.option_selected == "FONTE 40A" and price >= fonte40Premium:
                        continue;
            elif self.option_selected == "FONTE 60A":
                if listing_type == "Clássico" and price:
                    if self.option_selected == "FONTE 60A" and price >= fonte60Classico:
                        continue;
                elif listing_type == "Premium" and price:
                    if self.option_selected == "FONTE 60A" and price >= fonte60Premium:
                        continue;
                        
            elif self.option_selected == "FONTE 60A LITE":
                if listing_type == "Clássico" and price:
                    if self.option_selected == "FONTE 60A LITE" and price >= fonte60liteClassico:
                        continue;
                elif listing_type == "Premium" and price:
                    if self.option_selected == "FONTE 60A LITE" and price >= fonte60litePremium:
                        continue;
            elif self.option_selected == "FONTE 70A":
                if listing_type == "Clássico" and price:
                    if self.option_selected == "FONTE 70A" and price >= fonte70Classico:
                        continue;
                elif listing_type == "Premium" and price:
                    if self.option_selected == "FONTE 70A" and price >= fonte70Premium:
                        continue;
                        
                        
            elif self.option_selected == "FONTE 70A LITE":
                if listing_type == "Clássico" and price:
                    if self.option_selected == "FONTE 70A LITE" and price >= fonte70liteClassico:
                        continue;
                elif listing_type == "Premium" and price:
                    if self.option_selected == "FONTE 70A LITE" and price >= fonte70litePremium:
                        continue;
                        
            elif self.option_selected == "FONTE 90 BOB":
                
                if listing_type == "Clássico" and price:
                    if self.option_selected == "FONTE 90 BOB" and price >= fonte90bobClassico:
                        continue;
                elif listing_type == "Premium" and price:
                    if self.option_selected == "FONTE 90 BOB" and price >= fonte90bobPremium:
                        continue;
                        
            elif self.option_selected == "FONTE 120A":
                if listing_type == "Clássico" and price:
                    if self.option_selected == "FONTE 120A" and price >= fonte120Classico:
                        continue;
                elif listing_type == "Premium" and price:
                    if self.option_selected == "FONTE 120A" and price >= fonte120Premium:
                        continue;
                        
            elif self.option_selected == "FONTE 120A LITE":
                if listing_type == "Clássico" and price:
                    if self.option_selected == "FONTE 120A LITE" and price >= fonte120liteClassico:
                        continue;
                elif listing_type == "Premium" and price:
                    if self.option_selected == "FONTE 120A LITE" and price >= fonte120litePremium:
                        continue;
                        
            elif self.option_selected == "FONTE 120 BOB":
                if listing_type == "Clássico" and price:
                    if self.option_selected == "FONTE 120 BOB" and price >= fonte120bobClassico:
                        continue;
                elif listing_type == "Premium" and price:
                    if self.option_selected == "FONTE 120 BOB" and price >= fonte120bobPremium:
                        continue;
                        
            elif self.option_selected == "FONTE 200A":
                if listing_type == "Clássico" and price:
                    if self.option_selected == "FONTE 200A" and price >= fonte200Classico:
                        continue;
                elif listing_type == "Premium" and price:
                    if self.option_selected == "FONTE 200A" and price >= fonte200Premium:
                        continue;
                
            elif self.option_selected == "FONTE 200 MONO":
                if listing_type == "Clássico" and price:
                    if self.option_selected == "FONTE 200 MONO" and price >= fonte200monoClassico:
                        continue;
                elif listing_type == "Premium" and price:
                    if self.option_selected == "FONTE 200 MONO" and price >= fonte200monoPremium:
                        continue;
                        
            elif self.option_selected == "FONTE 200A LITE":
                if listing_type == "Clássico" and price:
                    if self.option_selected == "FONTE 200A LITE" and price >= fonte200liteClassico:
                        continue;
                elif listing_type == "Premium" and price:
                    if self.option_selected == "FONTE 200A LITE" and price >= fonte200litePremium:
                        continue;
                        
            elif self.option_selected == "FONTE 200 BOB":
                if listing_type == "Clássico" and price:
                    if self.option_selected == "FONTE 200 BOB" and price >= fonte200bobClassico:
                        continue;
                elif listing_type == "Premium" and price:
                    if self.option_selected == "FONTE 200 BOB" and price >= fonte200bobPremium:
                        continue;
            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': response.xpath('//h1[@class="ui-pdp-title"]'), 'loja': loja, 'price':price,'listing_type': listing_type, 'cupom': ""})
            
            if response.xpath('//li[@class="andes-pagination__button andes-pagination__button--next"]/a/@href'):
                next_page = response.xpath('//li[@class="andes-pagination__button andes-pagination__button--next"]/a/@href').get()
                yield scrapy.Request(url=next_page, callback=self.get_catalog)
        # //div[@class="ui-pdp-actions__container"]/input[@name="item_id"]]
        # ?pdp_filters=item_id:MLB5041771116
                       
    def get_cupom(self, response):
        url = response.meta['url']
        name = response.meta['name']
        price = response.meta['price']
        price_previsto = response.meta['parcelado']
        loja = response.meta['loja']
        tipo = response.meta['tipo']
        lugar = response.meta['lugar']
        cupom = response.meta['cupom']
        for i in response.xpath('/html/body/main/div/div/div[1]/div'):
            if i.xpath('.//div/div/div[1]/div[1]/div[2]/span/text()').get() != "Em produtos selecionados":
                doc.add_paragraph(f'Modelo: {self.option_selected_new}')
                doc.add_paragraph(f'URL: {url}')
                doc.add_paragraph(f'Nome: {name}')
                doc.add_paragraph(f'Preço: {price}')
                doc.add_paragraph(f'Preço Previsto: {price_previsto}')
                doc.add_paragraph(f'Loja: {loja}')
                doc.add_paragraph(f'Tipo: {tipo}')
                doc.add_paragraph(f'Lugar: {lugar}')
                doc.add_paragraph(f'Cupom: {cupom}')
                doc.add_paragraph("--------------------------------------------------------------------")
                doc.add_paragraph('')
                
                yield {
                    'url': url,
                    'name': name,
                    'price': price,
                    'price_previsto': price_previsto,
                    'loja': loja,
                    'tipo': tipo,
                    'lugar': lugar
                }
                doc.save(fr"dados/{self.option_selected_new}.docx")
        
    